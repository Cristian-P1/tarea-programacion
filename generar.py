import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    """Establece el color de fondo de una celda en una tabla."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading_with_spacing(doc, text, level, before=18, after=6):
    """Agrega un encabezado con control estricto de espaciado."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    return p

# 1. Inicializar Documento
doc = Document()

# Configurar márgenes estándar (1 pulgada)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# 2. Configurar Estilos Globales
style_normal = doc.styles['Normal']
font_normal = style_normal.font
font_normal.name = 'Calibri'
font_normal.size = Pt(11)
font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33) # Gris oscuro profesional

# Paleta de colores Corporativa (Azul tecnológico)
COLOR_PRIMARY = RGBColor(0x1F, 0x4E, 0x79) # Azul marino
COLOR_SECONDARY = RGBColor(0x2E, 0x75, 0xB6) # Azul medio
COLOR_HEX_BG = "F2F4F7" # Gris claro para tablas

for l in range(1, 4):
    style_h = doc.styles[f'Heading {l}']
    style_h.font.name = 'Calibri'
    style_h.font.bold = True
    if l == 1:
        style_h.font.size = Pt(18)
        style_h.font.color.rgb = COLOR_PRIMARY
    elif l == 2:
        style_h.font.size = Pt(14)
        style_h.font.color.rgb = COLOR_SECONDARY
    else:
        style_h.font.size = Pt(12)
        style_h.font.color.rgb = COLOR_SECONDARY

# 3. TÍTULO PRINCIPAL
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(24)
title_p.paragraph_format.space_after = Pt(12)
title_run = title_p.add_run("Historia de la Informática:\nPersonajes Influyentes y Evolución de los Lenguajes")
title_run.font.name = 'Calibri'
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_PRIMARY

# Subtítulo / Metadatos
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.paragraph_format.space_after = Pt(36)
meta_run = meta_p.add_run("Compilación Cronológica y Estructurada por Generaciones")
meta_run.font.italic = True
meta_run.font.size = Pt(11)
meta_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

# Introducción corta
intro_p = doc.add_paragraph("La historia de la informática se divide en dos grandes corrientes complementarias: los pioneros que crearon las bases tecnológicas y teóricas del hardware y la computación, y las generaciones de sistemas que definieron la evolución de los lenguajes de programación modernos. Este documento presenta un recorrido cronológico detallado por ambas vertientes.")
intro_p.paragraph_format.space_after = Pt(18)
intro_p.paragraph_format.line_spacing = 1.15

# =========================================================================
# PARTE 1: PERSONAJES INFLUYENTES
# =========================================================================
add_heading_with_spacing(doc, "Parte 1: Personajes Influyentes en la Computación", level=1, before=24, after=12)

personajes_data = [
    # Prehistoria
    ("Prehistoria y Siglo XIX", "Blaise Pascal", "1642", "Inventó la Pascalina, la primera calculadora mecánica construida a base de engranajes y ruedas."),
    ("Prehistoria y Siglo XIX", "Gottfried Leibniz", "1673", "Perfeccionó la calculadora mecánica e introdujo y documentó de forma moderna el sistema binario."),
    ("Prehistoria y Siglo XIX", "Charles Babbage", "1837", "Diseñó la Máquina Analítica, considerada el primer diseño conceptual de un computador de propósito general."),
    ("Prehistoria y Siglo XIX", "Ada Lovelace", "1843", "Creó el primer algoritmo destinado a ser procesado por la máquina de Babbage. Reconocida como la primera programadora de la historia."),
    
    # Primera Mitad Siglo XX
    ("Primera Mitad Siglo XX (1930 - 1950)", "Alan Turing", "1936", "Formuló el concepto de la Máquina de Turing. Estableció las bases fundamentales de la teoría de la computación y la IA."),
    ("Primera Mitad Siglo XX (1930 - 1950)", "Alonzo Church", "1936", "Desarrolló el Cálculo Lambda, una herramienta matemática que sentó las bases teóricas directas de la programación funcional."),
    ("Primera Mitad Siglo XX (1930 - 1950)", "Konrad Zuse", "1941", "Construyó la Z3, la cual se convirtió en la primera computadora digital, completamente binaria y Turing-completa del mundo."),
    ("Primera Mitad Siglo XX (1930 - 1950)", "John von Neumann", "1945", "Diseñó la Arquitectura de Von Neumann, vigente en los computadores actuales (unificación de CPU, memoria y periféricos)."),
    ("Primera Mitad Siglo XX (1930 - 1950)", "Grace Hopper", "1947", "Inventó el primer compilador de la historia (A-0) y popularizó el uso del término 'bug' para referirse a un error de software."),
    
    # Segunda Mitad Siglo XX
    ("Segunda Mitad Siglo XX (1950 - 1990)", "J. Bardeen, W. Brattain, W. Shockley", "1947", "Inventaron el transistor, componente fundamental que reemplazó a los tubos de vacío y miniaturizó la electrónica."),
    ("Segunda Mitad Siglo XX (1950 - 1990)", "Jack Kilby y Robert Noyce", "1958", "Desarrollaron el circuito integrado de manera independiente, permitiendo empaquetar miles de transistores en silicio."),
    ("Segunda Mitad Siglo XX (1950 - 1990)", "Dennis Ritchie y Ken Thompson", "1969-1972", "Crearon el sistema operativo Unix y el lenguaje de programación C, pilares del ecosistema de software moderno."),
    ("Segunda Mitad Siglo XX (1950 - 1990)", "Bill Gates y Paul Allen", "1975", "Fundaron Microsoft y lideraron la industria impulsando el uso del software empaquetado para computadoras personales."),
    ("Segunda Mitad Siglo XX (1950 - 1990)", "Steve Jobs y Steve Wozniak", "1976", "Fundaron Apple y masificaron el uso de la computadora personal con interfaces gráficas amigables para el usuario común."),
    ("Segunda Mitad Siglo XX (1950 - 1990)", "Tim Berners-Lee", "1989", "Inventó la World Wide Web (WWW), el protocolo HTTP y el lenguaje de marcado HTML, unificando la red global."),
    
    # Siglo XXI
    ("Siglo XXI (2000 - Presente)", "Linus Torvalds", "1991-Pres.", "Desarrolló el núcleo de Linux y posteriormente Git, transformando el desarrollo de software colaborativo y de código abierto."),
    ("Siglo XXI (2000 - Presente)", "L. Page, S. Brin, M. Zuckerberg", "2004-2007", "Revolucionaron la Web 2.0 y el manejo de datos masivos (Big Data) a través de motores de búsqueda avanzados y redes sociales."),
    ("Siglo XXI (2000 - Presente)", "Geoffrey Hinton y Yann LeCun", "2010s", "Padres del Deep Learning (Aprendizaje Profundo), técnica que desencadenó la revolución actual de la Inteligencia Artificial."),
    ("Siglo XXI (2000 - Presente)", "Sam Altman y Demis Hassabis", "Presente", "Líderes de las organizaciones OpenAI y Google DeepMind, pioneros en la búsqueda y despliegue de la Inteligencia Artificial General.")
]

current_era = ""
for era, personaje, fecha, hito in personajes_data:
    if era != current_era:
        current_era = era
        add_heading_with_spacing(doc, current_era, level=2, before=14, after=6)
    
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    r_name = p.add_run(f"{personaje} ({fecha}): ")
    r_name.bold = True
    r_name.font.color.rgb = COLOR_PRIMARY
    
    p.add_run(hito)

# =========================================================================
# PARTE 2: GENERACIONES Y LENGUAJES
# =========================================================================
add_heading_with_spacing(doc, "Parte 2: Generaciones de la Informática y Lenguajes de Programación", level=1, before=24, after=12)

generaciones_data = [
    {
        "titulo": "Primera Generación (1940 - 1956): Tubos de Vacío",
        "tecnologia": "Computadoras de dimensiones gigantescas que utilizaban tubos de vacío para procesar los circuitos eléctricos principales. Alto consumo de energía y propensas a fallas mecánicas térmicas.",
        "lenguajes": [
            ("Lenguaje Máquina", "1940s", "Consiste en secuencias binarias puras (0 y 1). Es el único que el hardware entiende nativamente. Altamente complejo de escribir y totalmente dependiente del diseño del microchip.")
        ]
    },
    {
        "titulo": "Segunda Generación (1956 - 1963): El Transistor",
        "tecnologia": "La introducción del transistor sustituyó los tubos de vacío. Permitió reducir drásticamente el tamaño físico, la generación de calor y el consumo energético, aumentando la fiabilidad.",
        "lenguajes": [
            ("Lenguaje Ensamblador", "1950s", "Introdujo códigos mnemónicos (como ADD, MOV) para representar instrucciones binarias, facilitando la programación humana cercana al hardware."),
            ("Fortran", "1957", "Desarrollado por John Backus en IBM. Es el primer lenguaje de alto nivel de la historia, optimizado para cálculos científicos y matemáticos complejos."),
            ("Lisp", "1958", "Creado por John McCarthy. Introdujo el procesamiento de listas y la computación simbólica. Se convirtió en el lenguaje pilar de la investigación en Inteligencia Artificial."),
            ("COBOL", "1959", "Diseñado bajo la influencia de Grace Hopper con un enfoque de lectura similar al inglés, enfocado al mundo empresarial, comercial y de administración pública.")
        ]
    },
    {
        "titulo": "Tercera Generación (1964 - 1971): Circuitos Integrados",
        "tecnologia": "Aparición de los microchips o circuitos integrados, que empaquetaban miles de transistores en una pequeña placa de silicio. Nace el concepto de sistema operativo para gestionar múltiples tareas.",
        "lenguajes": [
            ("BASIC", "1964", "Diseñado por John Kemeny y Thomas Kurtz para abrir la programación a estudiantes de disciplinas no científicas gracias a su sintaxis simplificada."),
            ("Pascal", "1970", "Creado por Niklaus Wirth como herramienta educativa para forzar la adopción de buenas prácticas mediante la programación estructurada."),
            ("Lenguaje C", "1972", "Desarrollado por Dennis Ritchie. Consiguió un balance perfecto entre control de bajo nivel y abstracción. Sigue siendo la base conceptual de casi todo el software moderno.")
        ]
    },
    {
        "titulo": "Cuarta Generación (1971 - 1989): Microprocesadores",
        "tecnologia": "Integración a gran escala (VLSI) que unificó todos los componentes de la CPU en un solo chip de silicio (el microprocesador). Esto propició el nacimiento y masificación de las PC domésticas.",
        "lenguajes": [
            ("SQL", "1974", "Desarrollado por IBM bajo un modelo relacional. Se convirtió en el estándar absoluto para la definición, manipulación y consulta de bases de datos centralizadas."),
            ("C++", "1983", "Creado por Bjarne Stroustrup como una extensión de C, incorporando el paradigma de Programación Orientada a Objetos (POO) sin sacrificar el rendimiento de bajo nivel."),
            ("Objective-C", "1984", "Híbrido de C y Smalltalk, seleccionado posteriormente por NeXT y Apple para ser la base estructural de sus sistemas operativos por varias décadas."),
            ("Perl", "1987", "Creado por Larry Wall. Destacó por sus potentes capacidades de procesamiento de texto y expresiones regulares, ideal para scripts en servidores de red Unix.")
        ]
    },
    {
        "titulo": "Quinta Generación (1989 - Presente): Conectividad, Nube e IA",
        "tecnologia": "Era caracterizada por la conectividad ubicua (Internet), computación móvil (smartphones), sistemas paralelos masivos, infraestructura en la nube y el despliegue a gran escala de redes neuronales e Inteligencia Artificial.",
        "lenguajes": [
            ("Python", "1991", "Diseñado por Guido van Rossum, priorizando la legibilidad extrema del código. Actualmente domina por completo las ciencias de datos, machine learning e IA."),
            ("Java", "1995", "Creado por Sun Microsystems con su máquina virtual (JVM) bajo la premisa de portabilidad total multiplataforma. Estándar indiscutible del backend empresarial."),
            ("JavaScript", "1995", "Desarrollado por Brendan Eich en Netscape. Evolucionó de ser un script ligero para navegadores a convertirse en el motor principal del desarrollo web moderno tanto en cliente como en servidor."),
            ("C#", "2000", "Lanzado por Microsoft dentro de su estrategia .NET. Combina la robustez de C++ con la facilidad y orientación de objetos limpia de Java."),
            ("Go / Rust", "2009-2010", "Creados por Google y Mozilla respectivamente. Go destaca por su concurrencia simplificada en la nube; Rust provee seguridad estricta en el manejo de memoria sin recolector de basura.")
        ]
    }
]

for gen in generaciones_data:
    add_heading_with_spacing(doc, gen["titulo"], level=2, before=16, after=6)
    
    # Detalle de la Tecnología
    p_tec = doc.add_paragraph()
    p_tec.paragraph_format.space_after = Pt(12)
    p_tec.paragraph_format.line_spacing = 1.15
    r_tec_lbl = p_tec.add_run("Tecnología base: ")
    r_tec_lbl.bold = True
    p_tec.add_run(gen["tecnologia"])
    
    # Crear Tabla para los Lenguajes de la Generación
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    
    # Configurar anchos de columna fijos
    col_widths = [Inches(1.5), Inches(0.8), Inches(4.2)]
    
    # Encabezado de la tabla
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Lenguaje", "Año", "Descripción e Impacto Histórico"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1F4E79") # Color primario hex
        # Estilo de texto de encabezado
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    # Agregar filas de datos
    for lang_name, lang_year, lang_desc in gen["lenguajes"]:
        row_cells = table.add_row().cells
        
        row_cells[0].text = lang_name
        row_cells[1].text = lang_year
        row_cells[2].text = lang_desc
        
        # Aplicar anchos y formato fino a las celdas de datos
        for i in range(3):
            row_cells[i].width = col_widths[i]
            p = row_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.05
            for run in p.runs:
                run.font.size = Pt(9.5)
            
            # Formato específico para la primera columna (Nombre del lenguaje)
            if i == 0:
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = COLOR_SECONDARY
                
        # Agregar un sombreado sutil alterno para legibilidad (opcional, aquí plano)
        for cell in row_cells:
            set_cell_background(cell, "FAFAFA")

    # Espacio después de la tabla
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(6)
    p_spacer.paragraph_format.space_after = Pt(6)

# 4. Guardar archivo final
doc.save("Historia_de_la_Computacion.docx")
print("¡Archivo 'Historia_de_la_Computacion.docx' creado con éxito!")
